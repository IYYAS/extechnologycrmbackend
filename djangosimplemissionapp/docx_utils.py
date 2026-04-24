from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def generate_activity_docx(activities, context):
    """
    Generates a Word (.docx) report for employee daily activities.
    """
    document = Document()
    
    # Title
    title = document.add_heading(context.get('title', 'Employee Activity Report'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Employee: {context.get('employee_name', 'All Employees')} | {context.get('date_range', '')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    
    document.add_paragraph() # Spacer
    
    # Summary Table
    if len(activities) > 0:
        total_activities = len(activities)
        actual_sum = sum(100 - (a.pending_work_percentage or 0) for a in activities)
        target_sum = sum(getattr(a, 'target_work_percentage', 0) for a in activities)
        avg_progress = actual_sum / total_activities
        avg_target = target_sum / total_activities
        efficiency = (actual_sum / target_sum * 100) if target_sum > 0 else 0
        
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        row = table.rows[0].cells
        row[0].text = "Avg. Actual Progress"
        row[1].text = f"{avg_progress:.1f}%"
        
        row = table.add_row().cells
        row[0].text = "Avg. Target Progress"
        row[1].text = f"{avg_target:.1f}%"
        
        row = table.add_row().cells
        row[0].text = "Progress Efficiency"
        row[1].text = f"{efficiency:.1f}%"
        
        document.add_paragraph() # Spacer

    # Main Data Table
    table = document.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Date', 'Employee', 'Project', 'Role', 'Team', 'Description', 'Status', 'Target', 'Progress']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        # Bold headers
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        
    for activity in activities:
        row_cells = table.add_row().cells
        
        # Date
        row_cells[0].text = activity.date.strftime('%Y-%m-%d')
        
        # Employee
        row_cells[1].text = activity.employee.username if activity.employee else "N/A"
        
        # Project
        p_obj = activity.project
        row_cells[2].text = (p_obj.name if p_obj else None) or "N/A"
        
        # Role
        row_cells[3].text = str(getattr(activity, 'role', None) or "N/A")
        
        # Team
        t_obj = activity.team
        row_cells[4].text = (t_obj.name if t_obj else None) or "N/A"
        
        # Description
        row_cells[5].text = str(activity.description or "")
        
        # Status
        row_cells[6].text = "Delayed" if activity.is_timeline_exceeded else "On Time"
        
        # Target
        row_cells[7].text = f"{getattr(activity, 'target_work_percentage', 0)}%"
        
        # Progress
        row_cells[8].text = f"{100 - (activity.pending_work_percentage or 0)}%"

    # Total Footer
    document.add_paragraph()
    footer = document.add_paragraph(f"Total Activities: {len(activities)}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    target_stream = io.BytesIO()
    document.save(target_stream)
    target_stream.seek(0)
    return target_stream
